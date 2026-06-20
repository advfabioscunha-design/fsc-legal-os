"""FSC LEGAL OS v4.0 — Backend FastAPI (api.seudominio.com.br)."""
import httpx
from fastapi import FastAPI, Request, HTTPException, Header, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from .core.config import get_settings
from .core.db import get_db
from .agentes import triagem, especialista, jurisprudencial, radar
from .agentes.orquestrador import mudar_estado, escalar_para_humano, TransicaoInvalida
from .integracoes import asaas, zapsign, whatsapp

app = FastAPI(title="FC Legal OS", version="4.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # restringir ao domínio do app em produção
    allow_methods=["*"], allow_headers=["*"],
)


@app.on_event("startup")
def _agendar_radar():
    """Liga o Radar Jurimétrico semanal (DataJud) dentro do container da API.
    Desligar com RADAR_AUTO=false. Roda 1x/semana (padrão: segunda 06:00 UTC)."""
    s = get_settings()
    if not s.radar_auto:
        return
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        from apscheduler.triggers.cron import CronTrigger
        sched = BackgroundScheduler(timezone="UTC")
        sched.add_job(
            radar.radar_semanal,
            CronTrigger(day_of_week=s.radar_dia_semana, hour=s.radar_hora, minute=0),
            id="radar_semanal", replace_existing=True, max_instances=1,
        )
        # Agente de Relacionamento: parabéns de aniversário (diário, 12:00 UTC ~ 09:00 BRT)
        try:
            from .agentes import relacionamento
            sched.add_job(
                relacionamento.parabenizar_aniversariantes,
                CronTrigger(hour=12, minute=0),
                id="aniversarios", replace_existing=True, max_instances=1,
            )
        except Exception as e:
            print(f"[aniversarios] job não agendado: {e}")
        sched.start()
        app.state.scheduler = sched
    except Exception as e:  # API sobe mesmo sem o scheduler
        print(f"[radar] scheduler não iniciado: {e}")


@app.get("/health")
def health():
    s = get_settings()
    return {"status": "ok", "ambiente": s.ambiente, "versao": "4.0"}


# ── Portal: novo lead e conversa ─────────────────────────────────
class NovoLead(BaseModel):
    nome: str
    contato: str
    relato: str
    cpf: str | None = None
    canal: str = "PORTAL"


@app.post("/api/v1/leads")
def novo_lead(body: NovoLead):
    from .core.cpf import cpf_valido
    if body.cpf and not cpf_valido(body.cpf):
        raise HTTPException(400, "CPF inválido — confira os números.")
    return triagem.criar_caso(body.nome, body.contato, body.relato, body.canal, cpf=body.cpf)


@app.get("/api/v1/validar-cpf")
def validar_cpf_endpoint(cpf: str):
    from .core.cpf import cpf_valido
    return {"cpf": cpf, "valido": cpf_valido(cpf)}


class Mensagem(BaseModel):
    conteudo: str
    canal: str = "PORTAL"


@app.post("/api/v1/casos/{caso_id}/mensagens")
def conversar(caso_id: str, body: Mensagem):
    return especialista.atender(caso_id, body.conteudo, body.canal)


# ── Serviço de Elaboração de Contrato ────────────────────────────
class ContratoInit(BaseModel):
    nome: str
    contato: str
    cpf: str | None = None
    canal: str = "PORTAL"


@app.post("/api/v1/contrato/iniciar")
def contrato_iniciar(body: ContratoInit):
    from .agentes import contrato
    return contrato.iniciar(body.nome, body.contato, body.canal, cpf=body.cpf)


@app.post("/api/v1/contrato/{caso_id}/mensagens")
def contrato_conversar(caso_id: str, body: Mensagem):
    from .agentes import contrato
    return contrato.atender(caso_id, body.conteudo, body.canal)


# ── CRM ──────────────────────────────────────────────────────────
@app.get("/api/v1/casos")
def listar_casos(estado: str | None = None, grupo: str | None = None, situacao: str = "ATIVO"):
    db = get_db()
    def montar(com_situacao: bool):
        q = db.table("casos").select("*, clientes(nome,whatsapp,email,origem)")
        if com_situacao and situacao and situacao != "TODOS":
            q = q.eq("situacao", situacao)
        if estado:
            q = q.eq("estado", estado)
        if grupo:
            q = q.eq("grupo", grupo)
        return q.order("atualizado_em", desc=True).limit(300).execute().data
    try:
        return montar(True)
    except Exception:
        return montar(False)  # coluna 'situacao' ainda não criada (migração pendente)


@app.get("/api/v1/casos/{caso_id}")
def detalhe_caso(caso_id: str):
    db = get_db()
    caso = db.table("casos").select("*, clientes(*)").eq("id", caso_id).single().execute().data
    caso["mensagens"] = db.table("mensagens").select("*").eq("caso_id", caso_id) \
                          .order("criado_em").execute().data
    caso["documentos"] = db.table("documentos").select("*").eq("caso_id", caso_id).execute().data
    return caso


@app.post("/api/v1/casos/{caso_id}/aprovar-protocolar")
def aprovar_e_protocolar(caso_id: str, tribunal: str = "EPROC_TJSC"):
    """Botão APROVAR E PROTOCOLAR do CRM → fila do RPA."""
    try:
        mudar_estado(caso_id, "APROVADO", motivo="Aprovado pelo advogado no CRM")
        mudar_estado(caso_id, "PROTOCOLO_RPA", motivo=f"Enfileirado p/ {tribunal}")
    except TransicaoInvalida as e:
        raise HTTPException(409, str(e))

    get_db().table("protocolos").insert(
        {"caso_id": caso_id, "tribunal": tribunal}
    ).execute()
    try:
        from workers.fila import enfileirar_protocolo
        enfileirar_protocolo(caso_id, tribunal)
    except Exception:
        pass  # worker pega da tabela mesmo sem Redis
    return {"ok": True, "fila": tribunal}


@app.post("/api/v1/casos/{caso_id}/escalar")
def escalar_manual(caso_id: str, motivo: str = "PEDIDO_CLIENTE", detalhe: str = ""):
    return escalar_para_humano(caso_id, motivo, detalhe)


# ── Montagem de Processo do Escritório (operação interna) ─────────
_GRUPOS_VALIDOS = {"BANCARIO", "IMOBILIARIO", "TRABALHISTA",
                   "PREVIDENCIARIO", "TRIBUTARIO", "CONSUMIDOR", "OUTROS"}
_FASES_VALIDAS = {"QUALIFICACAO", "PROPOSTA", "CONTRATO", "PAGAMENTO",
                  "COLETA_DOCS", "COLETA_PROVAS", "ANALISE", "PETICAO", "REVISAO",
                  "PROTOCOLO_RPA", "PROTOCOLADO"}


class CasoEscritorio(BaseModel):
    nome: str
    contato: str | None = None      # email ou whatsapp
    cpf: str | None = None
    grupo: str | None = None        # grupo_tese
    fase: str = "PETICAO"           # estado em que o processo entra
    descricao: str | None = None
    honorarios: str | None = None
    numero_processo: str | None = None


@app.post("/api/v1/casos/escritorio")
def criar_caso_escritorio(body: CasoEscritorio):
    """Cadastra um processo do escritório (cliente atendido por fora) já na
    fase em que ele se encontra, para entrar direto na esteira de produção."""
    from .core.db import registrar_evento
    if not body.contato or not body.contato.strip():
        raise HTTPException(400, "Informe o contato do cliente (e-mail ou WhatsApp) para acionamento.")
    db = get_db()
    fase = body.fase if body.fase in _FASES_VALIDAS else "PETICAO"
    grupo = body.grupo if body.grupo in _GRUPOS_VALIDOS else None

    cli: dict = {"nome": body.nome, "origem": "ESCRITORIO"}
    if body.contato:
        if "@" in body.contato:
            cli["email"] = body.contato
        else:
            cli["whatsapp"] = "".join(ch for ch in body.contato if ch.isdigit())
    if body.cpf:
        cli["cpf_cnpj"] = body.cpf
    cliente = db.table("clientes").insert(cli).execute().data[0]

    caso_row: dict = {
        "cliente_id": cliente["id"],
        "estado": fase,
        "relato_inicial": body.descricao or "Processo cadastrado pelo escritório",
    }
    if grupo:
        caso_row["grupo"] = grupo
    if body.honorarios:
        caso_row["honorarios_valor"] = body.honorarios
    if body.numero_processo:
        caso_row["numero_processo"] = body.numero_processo
    caso = db.table("casos").insert(caso_row).execute().data[0]
    registrar_evento(caso["id"], "CASO_ESCRITORIO_CRIADO",
                     {"fase": fase, "origem": "ESCRITORIO"})
    return {"ok": True, "caso_id": caso["id"], "estado": fase}


@app.get("/api/v1/pendencias")
def pendencias():
    """Casos escalados para humano (caixa de Intervenção Urgente), com a
    contagem de mensagens do cliente ainda não respondidas."""
    db = get_db()
    casos = db.table("casos").select("*, clientes(nome,whatsapp,email,origem)") \
        .eq("estado", "ESCALADO_HUMANO").order("atualizado_em", desc=True) \
        .limit(100).execute().data
    for c in casos:
        msgs = db.table("mensagens").select("autor").eq("caso_id", c["id"]) \
            .order("criado_em").execute().data
        nao_resp = 0
        for m in reversed(msgs):
            if m["autor"] == "CLIENTE":
                nao_resp += 1
            else:
                break
        c["mensagens_nao_respondidas"] = nao_resp
    return casos


# ── Gestão do caso (detalhe, edição, anexos, situação) ───────────
from datetime import datetime, timezone as _tz


class MotivoBody(BaseModel):
    motivo: str | None = None


class NotaBody(BaseModel):
    texto: str


class DocLink(BaseModel):
    nome: str | None = None
    url: str
    tipo: str | None = "LINK"


class EditarCaso(BaseModel):
    relato_inicial: str | None = None
    grupo: str | None = None
    honorarios: str | None = None
    numero_processo: str | None = None


def _set_situacao(caso_id: str, situacao: str, motivo: str | None, evento: str):
    from .core.db import registrar_evento
    get_db().table("casos").update({
        "situacao": situacao, "situacao_motivo": motivo,
        "atualizado_em": datetime.now(_tz.utc).isoformat(),
    }).eq("id", caso_id).execute()
    registrar_evento(caso_id, evento, {"motivo": motivo})
    return {"ok": True, "situacao": situacao}


@app.post("/api/v1/casos/{caso_id}/suspender")
def suspender_caso(caso_id: str, body: MotivoBody):
    return _set_situacao(caso_id, "SUSPENSO", body.motivo, "CASO_SUSPENSO")


@app.post("/api/v1/casos/{caso_id}/arquivar")
def arquivar_caso(caso_id: str, body: MotivoBody):
    return _set_situacao(caso_id, "ARQUIVADO", body.motivo, "CASO_ARQUIVADO")


@app.post("/api/v1/casos/{caso_id}/ativar")
def ativar_caso(caso_id: str):
    return _set_situacao(caso_id, "ATIVO", None, "CASO_ATIVADO")


@app.delete("/api/v1/casos/{caso_id}")
def excluir_caso(caso_id: str):
    db = get_db()
    # 1) BACKUP completo na lixeira ANTES de apagar (retenção 6 meses)
    try:
        caso = db.table("casos").select("*, clientes(*)").eq("id", caso_id).single().execute().data
        snap = {
            "caso": caso,
            "mensagens": db.table("mensagens").select("*").eq("caso_id", caso_id).execute().data,
            "documentos": db.table("documentos").select("*").eq("caso_id", caso_id).execute().data,
            "eventos": db.table("eventos").select("*").eq("caso_id", caso_id).execute().data,
        }
        rotulo = (caso.get("clientes") or {}).get("nome") or caso_id[:8]
        db.table("lixeira").insert({"caso_id": caso_id, "rotulo": rotulo, "dados": snap}).execute()
    except Exception:
        pass  # se a lixeira não existir ainda, não bloqueia a exclusão
    # 2) Apaga das tabelas operacionais
    for t in ("mensagens", "documentos", "eventos", "protocolos"):
        try:
            db.table(t).delete().eq("caso_id", caso_id).execute()
        except Exception:
            pass
    db.table("casos").delete().eq("id", caso_id).execute()
    return {"ok": True, "excluido": caso_id, "backup": "lixeira (6 meses)"}


@app.get("/api/v1/lixeira")
def listar_lixeira():
    return get_db().table("lixeira").select("id,caso_id,rotulo,excluido_em,expira_em") \
        .order("excluido_em", desc=True).limit(200).execute().data


@app.post("/api/v1/lixeira/{item_id}/restaurar")
def restaurar_lixeira(item_id: str):
    """Restaura um caso excluído a partir do backup da lixeira."""
    db = get_db()
    item = db.table("lixeira").select("*").eq("id", item_id).single().execute().data
    snap = item.get("dados") or {}
    caso = dict(snap.get("caso") or {})
    caso.pop("clientes", None)  # campo aninhado da query, não é coluna
    if not caso.get("id"):
        raise HTTPException(400, "Backup inválido.")
    db.table("casos").upsert(caso).execute()
    for m in snap.get("mensagens", []):
        m2 = {k: v for k, v in m.items() if k != "id"}
        try: db.table("mensagens").insert(m2).execute()
        except Exception: pass
    for d in snap.get("documentos", []):
        d2 = {k: v for k, v in d.items() if k != "id"}
        try: db.table("documentos").insert(d2).execute()
        except Exception: pass
    db.table("lixeira").delete().eq("id", item_id).execute()
    return {"ok": True, "restaurado": caso.get("id")}


@app.patch("/api/v1/casos/{caso_id}")
def editar_caso(caso_id: str, body: EditarCaso):
    campos = {k: v for k, v in body.model_dump().items() if v is not None}
    if "honorarios" in campos:
        campos["honorarios_valor"] = campos.pop("honorarios")
    if campos:
        campos["atualizado_em"] = datetime.now(_tz.utc).isoformat()
        get_db().table("casos").update(campos).eq("id", caso_id).execute()
    return {"ok": True}


@app.post("/api/v1/casos/{caso_id}/nota")
def add_nota(caso_id: str, body: NotaBody):
    get_db().table("mensagens").insert({
        "caso_id": caso_id, "canal": "CRM", "autor": "HUMANO", "conteudo": body.texto,
    }).execute()
    return {"ok": True}


@app.post("/api/v1/casos/{caso_id}/documentos")
def add_documento_link(caso_id: str, body: DocLink):
    get_db().table("documentos").insert({
        "caso_id": caso_id, "tipo": body.tipo or "LINK",
        "storage_path": body.url, "observacao": body.nome, "status": "RECEBIDO",
    }).execute()
    return {"ok": True}


@app.post("/api/v1/casos/{caso_id}/documentos/upload")
async def upload_documento(caso_id: str, arquivo: UploadFile = File(...)):
    import re, uuid
    s = get_settings()
    db = get_db()
    conteudo = await arquivo.read()
    nome_orig = arquivo.filename or "arquivo"
    base = nome_orig.rsplit("/", 1)[-1].rsplit("\\", 1)[-1]
    # storage não aceita espaços/acentos na chave -> troca por "_"
    safe = re.sub(r"[^A-Za-z0-9._-]", "_", base) or "arquivo"
    path = f"{caso_id}/{uuid.uuid4().hex}_{safe}"   # chave única e válida
    try:
        db.storage.from_(s.bucket_documentos).upload(
            path, conteudo,
            {"content-type": arquivo.content_type or "application/octet-stream", "upsert": "true"},
        )
    except Exception as e:
        raise HTTPException(500, f"Falha no upload de '{arquivo.filename}': {e}")
    db.table("documentos").insert({
        "caso_id": caso_id, "tipo": "UPLOAD", "storage_path": path,
        "observacao": arquivo.filename, "status": "RECEBIDO",
    }).execute()
    return {"ok": True, "path": path}


@app.get("/api/v1/documentos/{doc_id}/url")
def url_documento(doc_id: str):
    """Gera link temporário para abrir um documento salvo no storage privado."""
    s = get_settings()
    db = get_db()
    d = db.table("documentos").select("storage_path").eq("id", doc_id).single().execute().data
    sp = d.get("storage_path") or ""
    if sp.startswith("http"):
        return {"url": sp}
    try:
        res = db.storage.from_(s.bucket_documentos).create_signed_url(sp, 3600)
        url = res.get("signedURL") or res.get("signedUrl") or res.get("signed_url")
        return {"url": url}
    except Exception as e:
        raise HTTPException(500, f"Falha ao gerar link: {e}")


@app.delete("/api/v1/documentos/{doc_id}")
def excluir_documento(doc_id: str):
    """Remove o documento da pasta do cliente (storage) e do banco."""
    s = get_settings()
    db = get_db()
    d = db.table("documentos").select("storage_path").eq("id", doc_id).single().execute().data
    sp = (d or {}).get("storage_path") or ""
    if sp and not sp.startswith("http"):
        try:
            db.storage.from_(s.bucket_documentos).remove([sp])
        except Exception:
            pass
    db.table("documentos").delete().eq("id", doc_id).execute()
    return {"ok": True}


# ── SPRINT 1: Motor de Intimações e Processos (Escavador) ────────
class MonitorarBody(BaseModel):
    numero_processo: str | None = None
    oab: str | None = None
    tribunal: str | None = None
    caso_id: str | None = None


class StatusIntimacao(BaseModel):
    status: str


@app.post("/webhooks/escavador")
async def webhook_escavador(req: Request):
    """Recebe callbacks do Escavador. Insert rápido e idempotente, responde já."""
    from .integracoes import escavador
    try:
        payload = await req.json()
    except Exception:
        payload = {}
    return escavador.processar_webhook(payload)


@app.post("/api/v1/processos/monitorar")
def monitorar_processo(body: MonitorarBody):
    from .integracoes import escavador
    if not body.numero_processo and not body.oab:
        raise HTTPException(400, "Informe o número do processo ou a OAB.")
    return escavador.criar_monitoramento(body.numero_processo, body.oab, body.tribunal, body.caso_id)


@app.get("/api/v1/intimacoes")
def listar_intimacoes(status: str | None = None, tribunal: str | None = None,
                      q: str | None = None, limite: int = 200):
    qy = get_db().table("intimacoes").select("*")
    if status:
        qy = qy.eq("status", status)
    if tribunal:
        qy = qy.eq("tribunal", tribunal)
    if q:
        qy = qy.ilike("numero_processo", f"%{q}%")
    return qy.order("data_movimento", desc=True).limit(limite).execute().data


@app.patch("/api/v1/intimacoes/{intimacao_id}")
def atualizar_intimacao(intimacao_id: str, body: StatusIntimacao):
    if body.status not in ("A_RESOLVER", "RESOLVIDO", "PERDA_PRAZO"):
        raise HTTPException(400, "status inválido")
    get_db().table("intimacoes").update({"status": body.status, "lida": True}).eq("id", intimacao_id).execute()
    return {"ok": True}


@app.get("/api/v1/intimacoes/nao-lidas")
def intimacoes_nao_lidas():
    r = get_db().table("intimacoes").select("id").eq("lida", False).execute().data
    return {"total": len(r)}


@app.post("/api/v1/intimacoes/marcar-lidas")
def marcar_intimacoes_lidas():
    get_db().table("intimacoes").update({"lida": True}).eq("lida", False).execute()
    return {"ok": True}


@app.get("/api/v1/processos")
def listar_processos(tribunal: str | None = None, grupo: str | None = None,
                     estado: str | None = None, ano: str | None = None, limite: int = 200):
    rows = get_db().table("casos").select("*, clientes(nome,whatsapp,email)") \
        .order("atualizado_em", desc=True).limit(400).execute().data
    rows = [r for r in rows if r.get("numero_processo")]
    if tribunal:
        rows = [r for r in rows if r.get("tribunal") == tribunal]
    if grupo:
        rows = [r for r in rows if r.get("grupo") == grupo]
    if estado:
        rows = [r for r in rows if r.get("estado") == estado]
    if ano:
        rows = [r for r in rows if ano in (r.get("numero_processo") or "")]
    return rows[:limite]


@app.get("/api/v1/buscar")
def buscar_global(q: str):
    """Busca global da Top Bar: por nome, CPF/CNPJ ou número do processo."""
    db = get_db()
    termo = f"%{q}%"
    vistos: dict = {}
    try:
        for col in ("nome", "cpf_cnpj"):
            for c in db.table("clientes").select("id,nome,cpf_cnpj").ilike(col, termo).limit(15).execute().data:
                vistos[c["id"]] = c
    except Exception:
        pass
    casos = []
    if vistos:
        casos = db.table("casos").select("id,estado,grupo,numero_processo,cliente_id,clientes(nome,cpf_cnpj)") \
            .in_("cliente_id", list(vistos.keys())).limit(40).execute().data
    porproc = db.table("casos").select("id,estado,grupo,numero_processo,clientes(nome,cpf_cnpj)") \
        .ilike("numero_processo", termo).limit(20).execute().data
    # dedup por id de caso
    final = {c["id"]: c for c in (casos + porproc)}
    return {"casos": list(final.values())}


# ── SPRINT 2: Painel Administrativo e Financeiro ─────────────────
@app.get("/api/v1/admin/metricas")
def admin_metricas(periodo: str = "mes"):
    """KPIs do escritório a partir dos casos. periodo: dia|semana|mes|ano."""
    import re
    from datetime import datetime, timedelta, timezone as tz
    db = get_db()
    casos = db.table("casos").select(
        "estado,grupo,honorarios_valor,situacao,criado_em,pagamento_confirmado_em"
    ).limit(5000).execute().data

    agora = datetime.now(tz.utc)
    dias = {"dia": 1, "semana": 7, "mes": 30, "ano": 365}.get(periodo, 30)
    corte = agora - timedelta(days=dias)

    def _dt(s):
        try:
            return datetime.fromisoformat(str(s).replace("Z", "+00:00"))
        except Exception:
            return None

    no_periodo = [c for c in casos if (_dt(c.get("criado_em")) or agora) >= corte]

    por_nicho: dict = {}
    for c in no_periodo:
        g = c.get("grupo") or "OUTROS"
        por_nicho[g] = por_nicho.get(g, 0) + 1

    CONTRATACAO = {"LEAD", "QUALIFICACAO", "PROPOSTA", "CONTRATO", "PAGAMENTO"}
    PRODUCAO = {"COLETA_DOCS", "COLETA_PROVAS", "ANALISE", "PETICAO", "REVISAO", "APROVADO", "PROTOCOLO_RPA"}
    ATIVOS = {"PROTOCOLADO", "CONCLUIDO"}
    funil = {"contratacao": 0, "producao": 0, "ativos": 0}
    for c in casos:
        e = c.get("estado")
        if e in CONTRATACAO:
            funil["contratacao"] += 1
        elif e in PRODUCAO:
            funil["producao"] += 1
        elif e in ATIVOS:
            funil["ativos"] += 1

    def _valor(txt):
        t = str(txt or "")
        if "%" in t:
            return 0.0
        m = re.search(r"(\d{1,3}(?:\.\d{3})*(?:,\d{2})?|\d+(?:[.,]\d{2})?)", t)
        if not m:
            return 0.0
        try:
            return float(m.group(1).replace(".", "").replace(",", "."))
        except Exception:
            return 0.0

    recebido = a_receber = 0.0
    vals = []
    for c in casos:
        v = _valor(c.get("honorarios_valor"))
        if v > 0:
            vals.append(v)
        if c.get("pagamento_confirmado_em"):
            recebido += v
        elif c.get("estado") in (CONTRATACAO | PRODUCAO | ATIVOS) - {"LEAD", "QUALIFICACAO", "PROPOSTA"}:
            a_receber += v
    ticket = round(sum(vals) / len(vals), 2) if vals else 0.0

    return {
        "periodo": periodo,
        "total_casos": len(no_periodo),
        "por_nicho": por_nicho,
        "funil": funil,
        "financeiro": {
            "recebido": round(recebido, 2), "a_receber": round(a_receber, 2),
            "ticket_medio": ticket, "qtd_valores": len(vals),
        },
    }


# ── Módulo de Custos (entradas/saídas + folha) ───────────────────
class Lancamento(BaseModel):
    tipo: str                      # ENTRADA | SAIDA
    categoria: str | None = None
    descricao: str | None = None
    valor: float = 0
    data: str | None = None
    recorrente: bool = False


class FolhaItem(BaseModel):
    nome: str
    cargo: str | None = None
    salario: float = 0


@app.get("/api/v1/financeiro/lancamentos")
def listar_lancamentos(limite: int = 300):
    return get_db().table("fin_lancamentos").select("*").order("data", desc=True).limit(limite).execute().data


@app.post("/api/v1/financeiro/lancamentos")
def criar_lancamento(body: Lancamento):
    if body.tipo not in ("ENTRADA", "SAIDA"):
        raise HTTPException(400, "tipo inválido (use ENTRADA ou SAIDA)")
    row = {"tipo": body.tipo, "categoria": body.categoria, "descricao": body.descricao,
           "valor": body.valor, "recorrente": body.recorrente}
    if body.data:
        row["data"] = body.data
    get_db().table("fin_lancamentos").insert(row).execute()
    return {"ok": True}


@app.delete("/api/v1/financeiro/lancamentos/{lanc_id}")
def excluir_lancamento(lanc_id: str):
    get_db().table("fin_lancamentos").delete().eq("id", lanc_id).execute()
    return {"ok": True}


@app.get("/api/v1/financeiro/folha")
def listar_folha():
    return get_db().table("fin_folha").select("*").eq("ativo", True).order("criado_em").execute().data


@app.post("/api/v1/financeiro/folha")
def criar_folha(body: FolhaItem):
    get_db().table("fin_folha").insert({"nome": body.nome, "cargo": body.cargo, "salario": body.salario}).execute()
    return {"ok": True}


@app.delete("/api/v1/financeiro/folha/{folha_id}")
def excluir_folha(folha_id: str):
    get_db().table("fin_folha").delete().eq("id", folha_id).execute()
    return {"ok": True}


@app.get("/api/v1/financeiro/resumo")
def financeiro_resumo(periodo: str = "mes"):
    from datetime import datetime, timedelta, date, timezone as tz
    db = get_db()
    lanc = db.table("fin_lancamentos").select("*").limit(5000).execute().data
    folha = db.table("fin_folha").select("salario").eq("ativo", True).execute().data
    dias = {"dia": 1, "semana": 7, "mes": 30, "ano": 365}.get(periodo, 30)
    corte = datetime.now(tz.utc).date() - timedelta(days=dias)

    def _d(s):
        try:
            return date.fromisoformat(str(s)[:10])
        except Exception:
            return None

    per = [l for l in lanc if (_d(l.get("data")) or corte) >= corte]
    entradas = sum(float(l.get("valor") or 0) for l in per if l.get("tipo") == "ENTRADA")
    saidas = sum(float(l.get("valor") or 0) for l in per if l.get("tipo") == "SAIDA")
    folha_mensal = sum(float(f.get("salario") or 0) for f in folha)
    recorrente_mensal = sum(float(l.get("valor") or 0) for l in lanc if l.get("tipo") == "SAIDA" and l.get("recorrente"))
    projecao_anual = (recorrente_mensal + folha_mensal) * 12
    return {
        "periodo": periodo, "entradas": round(entradas, 2), "saidas": round(saidas, 2),
        "saldo": round(entradas - saidas, 2), "folha_mensal": round(folha_mensal, 2),
        "projecao_anual_custos": round(projecao_anual, 2),
    }


# ── SPRINT 3: Agenda e Agente Controlador ────────────────────────
class Membro(BaseModel):
    nome: str
    especialidades: list[str] = []
    lider: bool = False
    google_email: str | None = None


class Prazo(BaseModel):
    titulo: str
    descricao: str | None = None
    data: str
    responsavel_id: str | None = None
    especialidade: str | None = None
    caso_id: str | None = None


class DistribuirPrazo(BaseModel):
    titulo: str
    descricao: str | None = None
    data: str
    especialidade: str | None = None
    caso_id: str | None = None


@app.get("/api/v1/membros")
def listar_membros():
    return get_db().table("membros_equipe").select("*").eq("ativo", True).order("nome").execute().data


@app.post("/api/v1/membros")
def criar_membro(body: Membro):
    get_db().table("membros_equipe").insert({
        "nome": body.nome, "especialidades": body.especialidades,
        "lider": body.lider, "google_email": body.google_email,
    }).execute()
    return {"ok": True}


@app.delete("/api/v1/membros/{membro_id}")
def excluir_membro(membro_id: str):
    get_db().table("membros_equipe").update({"ativo": False}).eq("id", membro_id).execute()
    return {"ok": True}


@app.get("/api/v1/prazos")
def listar_prazos(responsavel_id: str | None = None, inicio: str | None = None, fim: str | None = None):
    qy = get_db().table("prazos").select("*, membros_equipe(nome,lider)")
    if responsavel_id:
        qy = qy.eq("responsavel_id", responsavel_id)
    if inicio:
        qy = qy.gte("data", inicio)
    if fim:
        qy = qy.lte("data", fim)
    return qy.order("data").limit(1000).execute().data


@app.post("/api/v1/prazos")
def criar_prazo(body: Prazo):
    get_db().table("prazos").insert({
        "titulo": body.titulo, "descricao": body.descricao, "data": body.data,
        "responsavel_id": body.responsavel_id, "especialidade": body.especialidade,
        "caso_id": body.caso_id, "origem": "MANUAL",
    }).execute()
    return {"ok": True}


@app.patch("/api/v1/prazos/{prazo_id}")
def concluir_prazo(prazo_id: str, status: str = "CONCLUIDO"):
    get_db().table("prazos").update({"status": status}).eq("id", prazo_id).execute()
    return {"ok": True}


@app.delete("/api/v1/prazos/{prazo_id}")
def excluir_prazo(prazo_id: str):
    get_db().table("prazos").delete().eq("id", prazo_id).execute()
    return {"ok": True}


@app.post("/api/v1/prazos/distribuir")
def distribuir_prazo(body: DistribuirPrazo):
    """Agente Controlador: distribui o prazo ao membro APTO (por especialidade)
    com MENOR carga de prazos abertos na semana (balanceamento de carga)."""
    from datetime import datetime, timedelta, date, timezone as tz
    from .core.db import registrar_evento
    db = get_db()
    membros = db.table("membros_equipe").select("*").eq("ativo", True).execute().data
    if not membros:
        raise HTTPException(400, "Cadastre membros na equipe antes de distribuir.")
    aptos = [m for m in membros if not body.especialidade or body.especialidade in (m.get("especialidades") or [])]
    if not aptos:
        aptos = membros  # ninguém com a tag: considera todos

    hoje = datetime.now(tz.utc).date()
    ini = hoje - timedelta(days=hoje.weekday())
    fim = ini + timedelta(days=7)
    carga: dict = {}
    for p in db.table("prazos").select("responsavel_id,data,status").eq("status", "ABERTO").execute().data:
        try:
            dd = date.fromisoformat(str(p.get("data"))[:10])
        except Exception:
            dd = None
        if dd and ini <= dd < fim and p.get("responsavel_id"):
            carga[p["responsavel_id"]] = carga.get(p["responsavel_id"], 0) + 1

    escolhido = min(aptos, key=lambda m: carga.get(m["id"], 0))
    row = db.table("prazos").insert({
        "titulo": body.titulo, "descricao": body.descricao, "data": body.data,
        "responsavel_id": escolhido["id"], "especialidade": body.especialidade,
        "caso_id": body.caso_id, "origem": "CONTROLADORIA",
    }).execute().data[0]
    registrar_evento(body.caso_id, "PRAZO_DISTRIBUIDO",
                     {"responsavel": escolhido["nome"], "carga_semana": carga.get(escolhido["id"], 0)})
    # notificação ao responsável: ativa quando o WhatsApp/Cloud API estiver ligado
    return {"ok": True, "responsavel": escolhido["nome"], "prazo_id": row["id"],
            "carga_semana": carga.get(escolhido["id"], 0)}


# ── SPRINT 4: Assistente CEO (interno) e Relacionamento ──────────
class AssistenteMsg(BaseModel):
    conteudo: str
    historico: list[dict] = []


@app.post("/api/v1/assistente")
def assistente_ceo(body: AssistenteMsg):
    from .agentes import ceo
    return ceo.responder(body.historico, body.conteudo)


@app.post("/api/v1/relacionamento/aniversarios")
def disparar_aniversarios():
    """Disparo manual das felicitações de aniversário (o automático roda diário)."""
    from .agentes import relacionamento
    return relacionamento.parabenizar_aniversariantes()


@app.post("/api/v1/casos/{caso_id}/iniciar")
def iniciar_caso_escritorio(caso_id: str):
    """Botão do CRM: inicia o processo do escritório na esteira (situação ATIVA)
    e registra o disparo para os agentes da esteira darem sequência."""
    from .core.db import registrar_evento
    get_db().table("casos").update({
        "situacao": "ATIVO", "atualizado_em": datetime.now(_tz.utc).isoformat(),
    }).eq("id", caso_id).execute()
    registrar_evento(caso_id, "PROCESSO_INICIADO_ESTEIRA", {"por": "ESCRITORIO"})
    return {"ok": True}


# ── Acionamento do cliente + aprovação humana da etapa ───────────
ORDEM_ESTEIRA = ["LEAD", "QUALIFICACAO", "PROPOSTA", "CONTRATO", "PAGAMENTO",
                 "COLETA_DOCS", "COLETA_PROVAS", "ANALISE", "PETICAO", "REVISAO",
                 "APROVADO", "PROTOCOLO_RPA", "PROTOCOLADO"]


class AcionarBody(BaseModel):
    solicitacao: str


@app.post("/api/v1/casos/{caso_id}/acionar-cliente")
def acionar_cliente(caso_id: str, body: AcionarBody):
    """Operador digita o que precisa do cliente; vai para o agente de triagem
    do WhatsApp coletar (mensagem/documento/assinatura/pagamento)."""
    from .core.db import registrar_evento
    db = get_db()
    texto = f"[SOLICITAÇÃO AO CLIENTE] {body.solicitacao}"
    db.table("mensagens").insert({
        "caso_id": caso_id, "canal": "WHATSAPP", "autor": "HUMANO", "conteudo": texto,
    }).execute()
    registrar_evento(caso_id, "SOLICITACAO_CLIENTE", {"texto": body.solicitacao})
    # processo sai da produção até o cliente complementar (visível na área do cliente)
    try:
        db.table("casos").update({
            "aguardando_cliente": True, "aguardando_desc": body.solicitacao,
            "atualizado_em": datetime.now(_tz.utc).isoformat(),
        }).eq("id", caso_id).execute()
    except Exception:
        pass
    # tenta enviar já pelo WhatsApp (quando o agente/Cloud API estiver ligado)
    enviado = False
    try:
        from .integracoes.whatsapp import enviar_para_cliente
        enviar_para_cliente(caso_id, body.solicitacao)
        enviado = True
    except Exception:
        pass
    return {"ok": True, "enviado_whatsapp": enviado}


@app.post("/api/v1/casos/{caso_id}/aprovar-etapa")
def aprovar_etapa(caso_id: str):
    """Aprovação humana após a conferência: avança o caso para a próxima etapa."""
    db = get_db()
    caso = db.table("casos").select("estado").eq("id", caso_id).single().execute().data
    atual = caso["estado"]
    if atual not in ORDEM_ESTEIRA:
        raise HTTPException(400, f"Estado '{atual}' não permite avanço manual.")
    i = ORDEM_ESTEIRA.index(atual)
    if i + 1 >= len(ORDEM_ESTEIRA):
        return {"ok": True, "info": "Caso já está na última etapa."}
    prox = ORDEM_ESTEIRA[i + 1]
    try:
        mudar_estado(caso_id, prox, motivo="Aprovação humana — avançar etapa")
    except TransicaoInvalida as e:
        raise HTTPException(409, str(e))
    return {"ok": True, "novo_estado": prox}


class MoverFase(BaseModel):
    fase: str


class AjusteBody(BaseModel):
    descricao: str


@app.post("/api/v1/casos/{caso_id}/retomar")
def retomar_producao(caso_id: str):
    """Cliente complementou o que faltava → volta para a produção."""
    from .core.db import registrar_evento
    try:
        get_db().table("casos").update({
            "aguardando_cliente": False, "aguardando_desc": None,
            "atualizado_em": datetime.now(_tz.utc).isoformat(),
        }).eq("id", caso_id).execute()
    except Exception:
        pass
    registrar_evento(caso_id, "CLIENTE_RETORNOU", {})
    return {"ok": True}


@app.post("/api/v1/casos/{caso_id}/mover-fase")
def mover_fase(caso_id: str, body: MoverFase):
    """Override manual da fase (ex.: reinserir após ação humana no ponto certo)."""
    from .core.db import registrar_evento
    if body.fase not in ORDEM_ESTEIRA and body.fase not in ("CONCLUIDO", "AGENDADO"):
        raise HTTPException(400, f"Fase '{body.fase}' inválida.")
    get_db().table("casos").update({
        "estado": body.fase, "atualizado_em": datetime.now(_tz.utc).isoformat(),
    }).eq("id", caso_id).execute()
    registrar_evento(caso_id, "FASE_MOVIDA_MANUAL", {"fase": body.fase})
    return {"ok": True, "novo_estado": body.fase}


@app.post("/api/v1/casos/{caso_id}/solicitar-ajuste")
def solicitar_ajuste(caso_id: str, body: AjusteBody):
    """Fase/opção inexistente → registra para a pasta de ajuste e implementação."""
    from .core.db import registrar_evento
    registrar_evento(caso_id, "SOLICITACAO_AJUSTE", {"descricao": body.descricao})
    return {"ok": True}


@app.get("/api/v1/casos/{caso_id}/download")
def baixar_dossie(caso_id: str):
    """Baixa o processo zipado: dossiê em Word (.docx) + documentos do storage
    + links de nuvem, para ação humana urgente."""
    import io, zipfile
    from fastapi.responses import Response
    from docx import Document
    s = get_settings()
    db = get_db()
    caso = db.table("casos").select("*, clientes(*)").eq("id", caso_id).single().execute().data
    msgs = db.table("mensagens").select("autor,conteudo,criado_em").eq("caso_id", caso_id).order("criado_em").execute().data
    docs = db.table("documentos").select("*").eq("caso_id", caso_id).execute().data
    cli = caso.get("clientes") or {}

    # Dossiê em Word
    doc = Document()
    doc.add_heading(f"Dossiê do Processo — {cli.get('nome', '')}", 0)
    doc.add_paragraph(f"Estado: {caso.get('estado')}  |  Grupo: {caso.get('grupo')}  |  Situação: {caso.get('situacao', 'ATIVO')}")
    doc.add_paragraph(f"Nº do processo: {caso.get('numero_processo') or '—'}")
    doc.add_paragraph(f"Contato: {cli.get('email') or cli.get('whatsapp') or '—'}  |  CPF/CNPJ: {cli.get('cpf_cnpj') or '—'}")
    doc.add_paragraph(f"Honorários: {caso.get('honorarios_valor') or '—'}")
    doc.add_heading("Relato / Informações coletadas", level=1)
    doc.add_paragraph(caso.get("relato_inicial") or "—")
    doc.add_heading("Histórico do atendimento", level=1)
    for m in msgs:
        doc.add_paragraph(f"[{m.get('autor')}] {m.get('conteudo')}")
    buf = io.BytesIO()
    doc.save(buf)
    docx_bytes = buf.getvalue()

    # Zip com dossiê + documentos + links
    zbuf = io.BytesIO()
    with zipfile.ZipFile(zbuf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("dossie.docx", docx_bytes)
        links = []
        for d in docs:
            sp = d.get("storage_path") or ""
            if sp.startswith("http"):
                links.append(f"{d.get('observacao') or d.get('tipo')}: {sp}")
            else:
                try:
                    data = db.storage.from_(s.bucket_documentos).download(sp)
                    z.writestr(f"documentos/{sp.split('/')[-1]}", data)
                except Exception:
                    pass
        if links:
            z.writestr("links_nuvem.txt", "\n".join(links))
    nome = (cli.get("nome") or "processo").replace(" ", "_")[:40]
    return Response(
        content=zbuf.getvalue(), media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="processo_{nome}.zip"'},
    )


@app.post("/api/v1/cerebro/processar")
def cerebro_processar():
    """Processa os acórdãos do bucket jurisprudencia/<GRUPO>[/<SUBNICHO>]/,
    cria/reforça teses e move os PDFs para _processados (automático)."""
    return jurisprudencial.processar_bucket()


@app.post("/api/v1/cerebro/radar-semanal")
def cerebro_radar_semanal(processar_pdfs: bool = True):
    """Radar Jurimétrico DataJud/CNJ: varre TJRO/TRT14 por grupo, agrega a
    jurimetria da semana, atualiza teses pelas íntegras e notifica o advogado.
    Disparado pelo scheduler semanal ou manualmente pelo CRM."""
    return radar.radar_semanal(processar_pdfs=processar_pdfs)


@app.get("/api/v1/radar")
def listar_radar(grupo: str | None = None, limite: int = 12):
    q = get_db().table("radar_jurimetrico").select("*")
    if grupo:
        q = q.eq("grupo", grupo)
    return q.order("semana_ref", desc=True).limit(limite).execute().data


# ── Área do cliente (autenticada via Supabase) ──────────────────
def _usuario_do_token(authorization: str | None) -> dict:
    """Valida o access_token do Supabase chamando /auth/v1/user
    (não precisa do JWT secret) e retorna o usuário {id, email}."""
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(401, "Token ausente")
    token = authorization.split(" ", 1)[1]
    s = get_settings()
    try:
        r = httpx.get(
            f"{s.supabase_url}/auth/v1/user",
            headers={"Authorization": f"Bearer {token}", "apikey": s.supabase_service_key},
            timeout=15,
        )
    except httpx.HTTPError:
        raise HTTPException(503, "Auth indisponível")
    if r.status_code != 200:
        raise HTTPException(401, "Token inválido")
    return r.json()


class RegistroEquipe(BaseModel):
    codigo: str


@app.post("/api/v1/equipe/registrar")
def equipe_registrar(body: RegistroEquipe, authorization: str | None = Header(default=None)):
    """Promove o usuário logado a OPERADOR se o código de acesso da equipe
    conferir. Permite cadastrar vários funcionários com segurança."""
    user = _usuario_do_token(authorization)
    s = get_settings()
    if not s.equipe_codigo or body.codigo.strip() != s.equipe_codigo:
        raise HTTPException(403, "Código de acesso inválido")
    get_db().table("perfis").upsert({
        "id": user["id"], "papel": "OPERADOR", "email": user.get("email"),
    }, on_conflict="id").execute()
    return {"ok": True, "papel": "OPERADOR"}


@app.get("/api/v1/cliente/meus-casos")
def cliente_meus_casos(authorization: str | None = Header(default=None)):
    """Casos do cliente logado — vinculados pelo e-mail do cadastro."""
    user = _usuario_do_token(authorization)
    email = user.get("email")
    if not email:
        return []
    db = get_db()
    cli = db.table("clientes").select("id").eq("email", email).maybe_single().execute().data
    if not cli:
        return []
    return db.table("casos").select("id,estado,grupo,subtipo") \
             .eq("cliente_id", cli["id"]).order("criado_em", desc=True).execute().data


@app.get("/api/v1/teses")
def listar_teses(grupo: str | None = None):
    q = get_db().table("teses").select("*").eq("overruled", False)
    if grupo:
        q = q.eq("grupo", grupo)
    return q.execute().data


# ── Webhooks (eventos confirmam, nunca o agente) ────────────────
@app.post("/webhooks/asaas")
async def webhook_asaas(req: Request):
    return asaas.processar_webhook(await req.json())


@app.post("/webhooks/zapsign")
async def webhook_zapsign(req: Request):
    return zapsign.processar_webhook(await req.json())


@app.post("/webhooks/whatsapp")
async def webhook_whatsapp(req: Request):
    return whatsapp.processar_webhook(await req.json())


@app.get("/webhooks/whatsapp")
def verificar_whatsapp(request: Request):
    """Verificação do webhook exigida pela Meta."""
    p = request.query_params
    if p.get("hub.verify_token") == "fsc-legal-os":
        return int(p.get("hub.challenge", 0))
    raise HTTPException(403)
